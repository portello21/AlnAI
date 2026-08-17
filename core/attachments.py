import csv
import hashlib
import io
import json
from pathlib import Path
from typing import Optional

from providers.ocr import extract_text_tesseract

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".log", ".py", ".ps1", ".yaml", ".yml", ".toml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
MAX_EXTRACTED_CHARS = 250_000


def calculate_file_sha256(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def decode_text_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _cap(text: str) -> str:
    return (text or "").strip()[:MAX_EXTRACTED_CHARS]


def extract_csv_text(file_bytes: bytes) -> str:
    raw = decode_text_bytes(file_bytes)
    try:
        return "\n".join(" | ".join(str(cell) for cell in row) for row in csv.reader(io.StringIO(raw))).strip()
    except Exception:
        return raw.strip()


def extract_json_text(file_bytes: bytes) -> str:
    raw = decode_text_bytes(file_bytes)
    try:
        return json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
    except Exception:
        return raw.strip()


def extract_pdf_text(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    parts = []
    for page in reader.pages[:200]:
        value = page.extract_text() or ""
        if value.strip():
            parts.append(value)
        if sum(len(x) for x in parts) >= MAX_EXTRACTED_CHARS:
            break
    return _cap("\n\n".join(parts))


def extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            if sum(len(x) for x in parts) >= MAX_EXTRACTED_CHARS:
                return _cap("\n".join(parts))
    return _cap("\n".join(parts))


def extract_xlsx_text(file_bytes: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets[:50]:
        parts.append(f"PLANILHA: {sheet.title}")
        for row in sheet.iter_rows(max_row=10_000, values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append(" | ".join(values))
            if sum(len(item) for item in parts) >= MAX_EXTRACTED_CHARS:
                return _cap("\n".join(parts))
    return _cap("\n".join(parts))


def extract_document_text(file_bytes: bytes, filename: str, mime_type: Optional[str] = None) -> dict:
    filename = Path(filename or "arquivo").name
    extension = Path(filename).suffix.lower()
    mime_type = (mime_type or "").lower()
    result = {
        "filename": filename,
        "extension": extension,
        "mime_type": mime_type,
        "file_hash": calculate_file_sha256(file_bytes),
        "text": "",
        "method": None,
        "success": False,
        "error": None,
    }
    try:
        if extension in TEXT_EXTENSIONS:
            result["text"], result["method"] = _cap(decode_text_bytes(file_bytes)), "text"
        elif extension == ".csv":
            result["text"], result["method"] = _cap(extract_csv_text(file_bytes)), "csv"
        elif extension == ".json":
            result["text"], result["method"] = _cap(extract_json_text(file_bytes)), "json"
        elif extension in IMAGE_EXTENSIONS or mime_type.startswith("image/"):
            ocr_text = extract_text_tesseract(file_bytes)
            if ocr_text.lower().startswith("erro no ocr"):
                raise RuntimeError(ocr_text)
            result["text"], result["method"] = _cap(ocr_text), "tesseract"
        elif extension == ".pdf" or mime_type == "application/pdf":
            result["text"], result["method"] = extract_pdf_text(file_bytes), "pypdf"
        elif extension == ".docx" or "wordprocessingml" in mime_type:
            result["text"], result["method"] = extract_docx_text(file_bytes), "python-docx"
        elif extension == ".xlsx" or "spreadsheetml" in mime_type:
            result["text"], result["method"] = extract_xlsx_text(file_bytes), "openpyxl"
        else:
            result["error"] = f"Formato não suportado: {extension or mime_type or 'desconhecido'}"
            return result
        if not result["text"].strip():
            result["error"] = "Nenhum texto foi extraído."
            return result
        result["success"] = True
        return result
    except Exception as exc:
        result["error"] = f"{type(exc).__name__}: {exc}"
        return result
